from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE


class CreateDocument:
    "Create Document for practicals"
    def __init__(self, template, subject=None, practicalno=None):
        self._document = Document(template)
        if subject is not None:
            self.set_subject(subject)
        if practicalno is not None:
            self.set_practical_no(practicalno)

    def set_font_style(self):
        """ set font style for whole document """
        self._charstyle = self._document.styles.add_style('SourceStyle', WD_STYLE_TYPE.PARAGRAPH)
        self._charstyle.font.name = 'Times New Roman'

    def check_style(self):
        """Check weather style for document is created or not """
        if not self.__dict__.get("_charstyle"):
            self.set_font_style()
            
    def set_font_size(self, size):
        """ Set font size for whole document """
        if self._charstyle:
            self._charstyle.font.size = Pt(size)

    def set_font_size(self, element, size):
        """ Set font size for the element givent """
        self.check_style()
        element.style.font.size = Pt(size)
            
    def set_subject(self, subject):
        """ set subject (header) """
        self._subject = self._document.paragraphs[2]
        self._subject.text = subject
        self._subject.alignment = 0 # 0 => left, 1 => center, 2 => right
        self.set_font_style()
        self._subject.style = 'SourceStyle'
        self.set_font_size(element = self._subject, size =18)

    def set_practical_no(self, practical_no):
        self._practical_no = self._document.paragraphs[3]
        self._practical_no.text = "Practical " + str(practical_no)
        self._practical_no.alignment = 0 # 0 => left, 1 => center, 2 => right
        self.check_style()
        self._practical_no.style = 'SourceStyle'
        self.set_font_size(element=self._practical_no,size=18)
        
    def set_aim(self, aim=''):
        self.check_style()
        doc_aim = self._document.add_paragraph("aim" + str(aim))
        self.set_font_size(element=doc_aim,size=12)
        
    def set_source(self, source=''):
        self.check_style()
        doc_source = self._document.add_paragraph(source)
        self.set_font_size(element= doc_source, size = 12)
        
    def set_output(self, output):
        self.check_style()
        try:
            self._document.add_picture(output, width=Inches(5))
        except:
            print("There was error adding picture !..")
            print("Make sure they are in same directory..")
    
    def set_aim_source_output(self, aim, source, output):
        print(aim)
        if aim is None or source is None or output is None:
            return
        self.set_aim(aim=aim)
        self.set_source(source=source)
        self.set_output(output=output)

        self._document.add_page_break()

    def save_doc(self, name=None):
        self._document.save(name)

#NEW_DOC = CreateDocument(template = "template.docx", subject = "Theory of Computation")
#NEW_DOC.set_aim_source_output(aim="practical to implement,.....",source=""" SOmelength text here would be adapted""", output='')
#NEW_DOC.save_doc('test.docx')
